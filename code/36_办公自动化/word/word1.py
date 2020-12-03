# 使用Python操作Word
from docx import Document
from docx.shared import Inches

# 空白文档
document = Document()
# document = Document('./test.docx') #打开原有文档
paragraph = document.add_paragraph('基于Python的办公自动化处理')
paragraph = document.add_paragraph('段落1:Word自动化')

document.add_section()  # 添加新章节
paragraph = document.add_paragraph('章节2-1')
document.add_section()  # 添加新章节
paragraph = document.add_paragraph('章节3-1')

sections = document.sections  # 返回所有章节引用的对象
print('章节数=', len(sections))  # 返回章节总数

# 得到指定section
section = sections[0]  # 返回指定章节的对象
section = document.sections[-1]  # 返回文档最后一个章节

# 得到section的信息，1英寸=914400像素
new_height = section.page_height
new_width = section.page_width
print('Section页面高={}, Section页面宽={}'.format(new_height, new_width))
print('左边距={}, 右边距={}, 上边距={}, 下边距={}'.format(section.left_margin, section.right_margin, section.top_margin,
                                              section.bottom_margin))
print('页眉距离={}, 页脚距离={}'.format(section.header_distance, section.footer_distance))

# 添加标题
document.add_heading('标题：楼宇进出人员名单', level=0)

# 添加表格
table = document.add_table(rows=3, cols=3)
# 返回行对象
row = table.rows[0]
# 给行对象的第n个单元格赋值
row.cells[0].text = '姓名'
row.cells[1].text = '温度'
row.cells[2].text = '时间'

# 返回表格的单元格对象，并赋值
cell = table.cell(1, 0)
cell.text = '张飞'
cell = table.cell(2, 0)
cell.text = '关羽'

# 添加图片--添加的图像以原始大小显示
document.add_picture('pic1.jpg')

# 保存文件
document.save('./test.docx')
